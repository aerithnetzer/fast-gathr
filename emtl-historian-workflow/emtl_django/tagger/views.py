from __future__ import annotations

import csv
import io
import json
from copy import deepcopy
from html import escape
from pathlib import Path
from typing import Any
from urllib.parse import urlencode
from uuid import uuid4

from django.db import DatabaseError, transaction
from django.http import Http404, HttpRequest, HttpResponse, HttpResponseNotAllowed, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.safestring import mark_safe
from django.utils.text import slugify
from django.utils.http import url_has_allowed_host_and_scheme

from .models import Clause, Document, NewIdProposal, ReviewNote, StageOutput
from .stage_config import (
    DEFAULT_TOGGLE_KEYS,
    STAGE_CHECKER_NOTES,
    STAGE_PROFILES,
    TOGGLE_KEYS,
    control_context,
    dependency_closure,
    dependency_warnings,
    missing_required_stages,
    ordered_stage_ids,
    pipeline_context,
    selected_stage_ids,
    stage_labels,
)
from .services.stage_runner import ChatbotStageRunner, build_event_candidate_packages
from .services.contracts import ProviderLabel
from .services.clause_persistence import replace_document_clauses
from .services.entity_persistence import persist_entity_registry_attempt
from .services.entity_review_handoff import ENTITY_RECORD_TYPES, propose_entity
from .services.export_contract import build_workflow_export
from .services.workflow_orchestrator import build_orchestration_plan


FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
SAMPLE_DOC_ID = "E-133-2-211_dep01"
EVENT_HEADWORD_REVIEW_STATE_LABELS = {
    "llm_selected_candidate": "Suggestion ready",
    "needs_historian_review": "Needs review",
    "llm_none_fit": "No fit found",
    "proposed_new_headword": "New headword proposed",
    "accepted_llm_choice": "Accepted",
    "accepted_alternate_candidate": "Accepted alternate",
    "rejected_all_candidates": "No suitable candidate recorded",
    "proposed_new_headword_pending_review": "Proposal pending review",
    "held_for_later": "Held for later",
}
EVENT_HEADWORD_REVIEW_STATE_CLASSES = {
    "llm_selected_candidate": "selected",
    "needs_historian_review": "review",
    "llm_none_fit": "none-fit",
    "proposed_new_headword": "proposal",
    "accepted_llm_choice": "accepted",
    "accepted_alternate_candidate": "accepted-alt",
    "rejected_all_candidates": "rejected-all",
    "proposed_new_headword_pending_review": "proposal",
    "held_for_later": "held",
}
EVENT_HEADWORD_REVIEW_DECISION_LABELS = {
    "choose_candidate": "Suggested headword",
    "needs_historian_review": "Needs historian review",
    "none_of_these_fit": "No suitable candidate found",
    "propose_new_headword": "Provisional new headword suggested",
}
EVENT_HEADWORD_REVIEW_DECISION_CLASSES = {
    "choose_candidate": "selected",
    "needs_historian_review": "review",
    "none_of_these_fit": "none-fit",
    "propose_new_headword": "proposal",
    "accepted_llm_choice": "accepted",
    "accepted_alternate_candidate": "accepted-alt",
    "rejected_all_candidates": "rejected-all",
    "proposed_new_headword_pending_review": "proposal",
    "held_for_later": "held",
}
EVENT_HEADWORD_REVIEW_RECORDED_STATE_LABELS = {
    "accepted_llm_choice": "Accepted",
    "accepted_alternate_candidate": "Accepted alternate",
    "rejected_all_candidates": "No suitable candidate recorded",
    "proposed_new_headword_pending_review": "Proposal pending review",
    "held_for_later": "Held for later",
}
EVENT_HEADWORD_REVIEW_MUTABLE_INITIAL_STATES = {
    "llm_selected_candidate",
    "needs_historian_review",
    "llm_none_fit",
    "proposed_new_headword",
}
EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY = "event_headword_review_local_decisions"
EVENT_HEADWORD_REVIEW_SESSION_NOTICE_KEY = "event_headword_review_notice"


def _load_json_fixture(filename: str) -> dict[str, Any]:
    path = FIXTURE_DIR / filename
    return json.loads(path.read_text(encoding="utf-8"))


def _load_event_headword_review_fixture() -> dict[str, Any]:
    return _load_json_fixture("event_headword_review_sample.json")


def _is_sample_document(document: Document) -> bool:
    return (
        document.doc_id == SAMPLE_DOC_ID
        and (document.metadata or {}).get("workflow_source") == "sample_document"
    )


def _safe_next_url(request: HttpRequest) -> str:
    fallback = reverse("tagger:workbench")
    next_url = str(request.POST.get("next") or request.GET.get("next") or fallback)
    if url_has_allowed_host_and_scheme(next_url, allowed_hosts={request.get_host()}):
        return next_url
    return fallback


def _set_home_notice(request: HttpRequest, message: str) -> None:
    request.session["home_notice"] = message


def _workflow_started(request: HttpRequest) -> bool:
    return bool(request.session.get("workflow_started"))


def _clear_active_workflow(request: HttpRequest) -> None:
    request.session.pop("active_document_pk", None)
    request.session.pop("workflow_started", None)
    request.session.pop("review_notes", None)


def _selected_toggles(request: HttpRequest) -> set[str]:
    selected = {key for key in TOGGLE_KEYS if request.GET.get(key) == "on"}
    if request.GET.get("options_applied") == "1":
        return selected
    return set(DEFAULT_TOGGLE_KEYS)


def _toggle_query(selected_toggles: set[str]) -> str:
    parts = ["options_applied=1"]
    for key in TOGGLE_KEYS:
        if key in selected_toggles:
            parts.append(f"{key}=on")
    return "&".join(parts)


def _workbench_url(
    *,
    selected_toggles: set[str],
    selected_clause_id: str,
    stage_id: str = "",
) -> str:
    params: list[tuple[str, str]] = [
        ("clause", selected_clause_id),
        ("options_applied", "1"),
    ]
    for key in TOGGLE_KEYS:
        if key in selected_toggles:
            params.append((key, "on"))
    if stage_id:
        params.append(("stage", stage_id))
    return f"{reverse('tagger:workbench')}?{urlencode(params)}"


def _status_from_fixture_label(label: str) -> str:
    normalized = str(label or "").lower()
    if "approved" in normalized:
        return NewIdProposal.Status.APPROVED
    if "reject" in normalized:
        return NewIdProposal.Status.REJECTED
    if "edit" in normalized:
        return NewIdProposal.Status.NEEDS_EDIT
    return NewIdProposal.Status.PENDING


def _seed_proposals_for_document(document: Document, outputs: dict[str, Any], *, reset: bool = False) -> None:
    for row in outputs.get("new_id_proposals", []):
        source_clause = Clause.objects.filter(
            document=document,
            clause_id=str(row.get("source_clause") or ""),
        ).first()
        proposal, created = NewIdProposal.objects.get_or_create(
            document=document,
            proposed_id=str(row.get("proposed_id") or "").strip(),
            defaults={
                "record_type": str(row.get("type") or "").strip(),
                "headword": str(row.get("headword") or "").strip(),
                "evidence_form": str(row.get("evidence_form") or "").strip(),
                "source_clause": source_clause,
                "status": _status_from_fixture_label(str(row.get("status") or "")),
                "payload": {
                    "source": "sample_output_placeholder",
                    "source_clause_label": str(row.get("source_clause") or ""),
                    "fixture_status": str(row.get("status") or ""),
                },
            },
        )
        if reset and not created:
            proposal.record_type = str(row.get("type") or "").strip()
            proposal.headword = str(row.get("headword") or "").strip()
            proposal.evidence_form = str(row.get("evidence_form") or "").strip()
            proposal.source_clause = source_clause
            proposal.status = _status_from_fixture_label(str(row.get("status") or ""))
            proposal.reviewer_note = ""
            proposal.payload = {
                "source": "sample_output_placeholder",
                "source_clause_label": str(row.get("source_clause") or ""),
                "fixture_status": str(row.get("status") or ""),
            }
            proposal.save(
                update_fields=[
                    "record_type",
                    "headword",
                    "evidence_form",
                    "source_clause",
                    "status",
                    "reviewer_note",
                    "payload",
                    "updated_at",
                ]
            )


def _stage_fixture_payload(
    stage_id: str,
    document: Document,
    outputs: dict[str, Any],
) -> tuple[dict[str, Any], str]:
    if stage_id == "summary_keywords":
        payload = {
            "summary": outputs.get("summary", {}),
            "keywords": outputs.get("keywords", {}),
        }
        raw_output = (
            f"{payload['summary'].get('text', '')}\n\n"
            f"Keywords: {', '.join(payload['keywords'].get('items', []))}"
        ).strip()
        return payload, raw_output
    if stage_id == "entity_registry":
        payload = {
            "entity_registry": outputs.get("entity_registry", []),
            "annotations": outputs.get("annotations", []),
            "new_id_proposals": outputs.get("new_id_proposals", []),
        }
        return payload, outputs.get("raw_output_excerpt", "")
    if stage_id == "clause_parser":
        clauses = [
            {
                "clause_id": clause.clause_id,
                "sequence": clause.sequence,
                "text": clause.text,
            }
            for clause in document.clauses.all()
        ]
        return {
            "clauses": clauses,
            "notice": "The current prototype represents Clause Parser output through the loaded clause list.",
        }, "Sample clause parser output represented by the current local clause list."
    if stage_id == "occurrences_registry":
        payload = {
            "occurrences_registry": outputs.get("events", []),
        }
        if _is_sample_document(document):
            payload["event_headword_review"] = _load_event_headword_review_fixture()
        return payload, outputs.get("raw_output_excerpt", "")
    if stage_id == "tag_assembler":
        return {
            "placeholder": outputs.get("tag_assembler_placeholder", ""),
            "full_tagset_placeholder": outputs.get("full_tagset_placeholder", ""),
        }, outputs.get("tag_assembler_placeholder", "")
    if stage_id == "key_narrative":
        return {
            "future_notice": outputs.get("key_narrative_future", ""),
        }, outputs.get("key_narrative_future", "")
    return {}, ""


def _event_headword_review_has_local_actions(payload: dict[str, Any]) -> bool:
    items = payload.get("event_review_items") or []
    return any(
        item.get("review_decision") or item.get("audit_log") or item.get("review_state") not in EVENT_HEADWORD_REVIEW_MUTABLE_INITIAL_STATES
        for item in items
        if isinstance(item, dict)
    )


def _seed_stage_outputs_for_document(
    document: Document,
    outputs: dict[str, Any],
    *,
    preserve_event_headword_review: bool = True,
    force_fixture: bool = False,
) -> None:
    valid_statuses = {choice[0] for choice in StageOutput.Status.choices}
    for stage_id, profile in STAGE_PROFILES.items():
        payload, raw_output = _stage_fixture_payload(stage_id, document, outputs)
        initial_status = StageOutput.Status.BLOCKED if profile.future else StageOutput.Status.LOADED
        stage_output, created = StageOutput.objects.get_or_create(
            document=document,
            stage=stage_id,
            defaults={
                "status": initial_status,
                "display_title": profile.label,
                "payload": payload,
                "raw_output": raw_output,
                "provenance": {
                    "source": "sample_output_placeholder",
                    "fixture_file": "sample_outputs.json",
                    "loaded_at": timezone.now().isoformat(),
                    "real_chatbot_execution": False,
                    "postgresql_commit": False,
                },
            },
        )
        if (
            not created
            and not force_fixture
            and (stage_output.provenance or {}).get("source") != "sample_output_placeholder"
        ):
            continue
        update_fields = ["display_title", "payload", "raw_output", "provenance", "updated_at"]
        if (
            preserve_event_headword_review
            and stage_id == "occurrences_registry"
            and isinstance(stage_output.payload, dict)
            and isinstance(stage_output.payload.get("event_headword_review"), dict)
            and _event_headword_review_has_local_actions(stage_output.payload["event_headword_review"])
        ):
            payload["event_headword_review"] = stage_output.payload["event_headword_review"]
        stage_output.display_title = profile.label
        stage_output.payload = payload
        stage_output.raw_output = raw_output
        stage_output.provenance = {
            **(stage_output.provenance or {}),
            "source": "sample_output_placeholder",
            "fixture_file": "sample_outputs.json",
            "real_chatbot_execution": False,
            "postgresql_commit": False,
        }
        if not created and stage_output.status not in valid_statuses:
            stage_output.status = initial_status
            update_fields.append("status")
        stage_output.save(update_fields=update_fields)


def _reset_stage_outputs_for_document(document: Document, outputs: dict[str, Any] | None = None) -> None:
    outputs = outputs or _load_json_fixture("sample_outputs.json")
    _seed_stage_outputs_for_document(
        document,
        outputs,
        preserve_event_headword_review=False,
        force_fixture=True,
    )
    for stage_output in document.stage_outputs.all():
        profile = STAGE_PROFILES.get(stage_output.stage)
        stage_output.status = StageOutput.Status.BLOCKED if profile and profile.future else StageOutput.Status.LOADED
        provenance = dict(stage_output.provenance or {})
        provenance.update(
            {
                "source": "sample_output_placeholder",
                "fixture_file": "sample_outputs.json",
                "workflow_reset_at": timezone.now().isoformat(),
                "real_chatbot_execution": False,
                "postgresql_commit": False,
            }
        )
        for key in (
            "checked_at",
            "continued_at",
            "last_rerun_request_at",
            "passed_forward_locally",
            "blocked_reason",
        ):
            provenance.pop(key, None)
        stage_output.provenance = provenance
        stage_output.save(update_fields=["status", "provenance", "updated_at"])


def _ensure_sample_document(*, reset_review_state: bool = False) -> Document:
    document_fixture = _load_json_fixture("sample_document.json")
    outputs = _load_json_fixture("sample_outputs.json")
    metadata = dict(document_fixture.get("metadata") or {})
    fixture_clauses = list(document_fixture.get("clauses", []))
    working_source_text = _join_working_text_segments(
        [str(clause.get("text") or "") for clause in fixture_clauses]
    )
    document, _created = Document.objects.update_or_create(
        doc_id=SAMPLE_DOC_ID,
        defaults={
            "archival_reference": str(metadata.get("archival_reference") or ""),
            "title": str(metadata.get("title") or "Sample document"),
            "document_type": str(metadata.get("document_type") or ""),
            "normalized_date": str(metadata.get("normalized_date") or ""),
            "source_file": "sample_document.json",
            "metadata": {
                **metadata,
                "workflow_source": "sample_document",
                "placeholder_outputs": True,
                "working_source_text": working_source_text,
                "working_source_text_source": "sample_document_fixture",
                "working_source_paragraph_count": _paragraph_count(working_source_text),
            },
        },
    )
    for index, clause in enumerate(fixture_clauses, start=1):
        Clause.objects.update_or_create(
            document=document,
            clause_id=str(clause.get("clause_id") or index).zfill(4),
            defaults={
                "text": str(clause.get("text") or ""),
                "sequence": index,
                "start_char": clause.get("start_char"),
                "end_char": clause.get("end_char"),
            },
        )
    _seed_proposals_for_document(document, outputs, reset=reset_review_state)
    if reset_review_state:
        _reset_stage_outputs_for_document(document, outputs)
    else:
        _seed_stage_outputs_for_document(document, outputs)
    return document


def _split_working_text_segments(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return []
    segments: list[str] = []
    current: list[str] = []
    for line in normalized.split("\n"):
        if line.strip():
            current.append(line.rstrip())
        elif current:
            segments.append("\n".join(current).strip())
            current = []
    if current:
        segments.append("\n".join(current).strip())
    return [segment for segment in segments if segment]


def _normalize_working_source_text(text: str) -> str:
    return str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def _join_working_text_segments(segments: list[str]) -> str:
    cleaned = [_normalize_working_source_text(segment) for segment in segments]
    return "\n\n".join(segment for segment in cleaned if segment)


def _paragraph_count(text: str) -> int:
    normalized = _normalize_working_source_text(text)
    if not normalized:
        return 0
    return len([block for block in normalized.split("\n\n") if block.strip()])


def _docx_run_text(run: Any) -> str:
    text = str(getattr(run, "text", "") or "")
    if not text or not getattr(getattr(run, "font", None), "superscript", False):
        return text
    core = text.strip()
    if not core:
        return text
    leading = text[: len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()) :]
    marked = f"^{core}" if len(core) <= 3 and " " not in core else f"^({core})"
    return f"{leading}{marked}{trailing}"


def _docx_paragraph_text(paragraph: Any) -> str:
    if getattr(paragraph, "runs", None):
        text = "".join(_docx_run_text(run) for run in paragraph.runs)
    else:
        text = str(getattr(paragraph, "text", "") or "")
    return " ".join(text.split()) if "\n" not in text else text.strip()


def _append_unique_segment(segments: list[str], text: str, seen: set[str]) -> None:
    normalized = text.strip()
    if not normalized:
        return
    key = " ".join(normalized.split()).lower()
    if key in seen:
        return
    seen.add(key)
    segments.append(normalized)


def _extract_docx_segments(data: bytes) -> list[str]:
    from docx import Document as DocxReader

    docx_file = DocxReader(io.BytesIO(data))
    segments: list[str] = []
    seen: set[str] = set()

    for section in docx_file.sections:
        for paragraph in section.header.paragraphs:
            _append_unique_segment(segments, _docx_paragraph_text(paragraph), seen)

    for paragraph in docx_file.paragraphs:
        _append_unique_segment(segments, _docx_paragraph_text(paragraph), seen)

    for table in docx_file.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            if cells:
                _append_unique_segment(segments, " | ".join(cells), seen)

    for section in docx_file.sections:
        for paragraph in section.footer.paragraphs:
            _append_unique_segment(segments, _docx_paragraph_text(paragraph), seen)

    return segments


def _upload_fallback_segments(filename: str, message: str) -> dict[str, Any]:
    return {
        "segments": [message],
        "status": "fallback",
        "note": message,
        "format": Path(filename).suffix.lower().lstrip(".") or "unknown",
    }


def _extract_upload_working_text(uploaded_file: Any) -> dict[str, Any]:
    filename = str(getattr(uploaded_file, "name", "") or "")
    suffix = Path(filename).suffix.lower()
    data = uploaded_file.read()

    if suffix == ".docx":
        try:
            segments = _extract_docx_segments(data)
        except ImportError:
            return _upload_fallback_segments(
                filename,
                (
                    f"Uploaded file captured: {filename}. DOCX text extraction requires "
                    "python-docx, which is not installed in this environment."
                ),
            )
        except Exception:
            return _upload_fallback_segments(
                filename,
                (
                    f"Uploaded file captured: {filename}. DOCX text extraction could not "
                    "read this file, so sample chatbot outputs are used as placeholder returns."
                ),
            )
        if not segments:
            return _upload_fallback_segments(
                filename,
                (
                    f"Uploaded file captured: {filename}. No readable DOCX body text was found."
                ),
            )
        return {
            "segments": segments,
            "status": "extracted",
            "note": (
                "Plain-text DOCX extraction preserved paragraph breaks where possible. "
                "Word formatting is simplified; superscript runs are marked with ^."
            ),
            "format": "docx",
        }

    if suffix in {".txt", ".md", ".csv", ".xml", ".html", ".htm"}:
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                text = data.decode(encoding)
                segments = _split_working_text_segments(text)
                return {
                    "segments": segments or ["Uploaded text file was empty."],
                    "status": "extracted" if segments else "fallback",
                    "note": "Plain-text upload displayed directly.",
                    "format": suffix.lstrip(".") or "text",
                }
            except UnicodeDecodeError:
                continue
        return _upload_fallback_segments(
            filename,
            f"Uploaded file captured: {filename}. A readable text preview could not be created.",
        )

    if suffix == ".pdf":
        return _upload_fallback_segments(
            filename,
            (
                f"Uploaded file captured: {filename}. PDF text extraction is not connected "
                "in this prototype; OCR and full PDF viewing are not implemented."
            ),
        )

    return _upload_fallback_segments(
        filename,
        (
            f"Uploaded file captured: {filename}. Full text extraction is not connected "
            "for this file type, so sample chatbot outputs are used as placeholder returns."
        ),
    )


def _create_uploaded_document(uploaded_file: Any) -> Document:
    now = timezone.now()
    filename = str(getattr(uploaded_file, "name", "") or "uploaded-document")
    base = slugify(Path(filename).stem)[:48] or "uploaded-document"
    doc_id = f"upload-{now:%Y%m%d%H%M%S%f}-{base}"
    extraction = _extract_upload_working_text(uploaded_file)
    segments = list(extraction.get("segments") or ["Uploaded file captured, but no working text was extracted."])
    working_source_text = _join_working_text_segments([str(segment) for segment in segments])
    document = Document.objects.create(
        doc_id=doc_id,
        archival_reference="",
        title=filename,
        document_type="uploaded document",
        normalized_date="",
        source_file=filename,
        metadata={
            "doc_id": doc_id,
            "title": filename,
            "document_title": filename,
            "document_type": "uploaded document",
            "uploaded_at": now.isoformat(),
            "workflow_source": "uploaded_document",
            "placeholder_outputs": False,
            "live_workflow_enabled": True,
            "upload_extraction_status": extraction.get("status", ""),
            "upload_extraction_note": extraction.get("note", ""),
            "upload_text_format": extraction.get("format", ""),
            "upload_segment_count": len(segments),
            "working_source_text": working_source_text,
            "working_source_text_source": "plain_text_upload_extraction",
            "working_source_paragraph_count": _paragraph_count(working_source_text),
        },
    )
    for index, segment in enumerate(segments, start=1):
        Clause.objects.create(
            document=document,
            clause_id=str(index).zfill(4),
            sequence=index,
            text=str(segment),
        )
    _ensure_live_stage_outputs(document)
    return document


def _ensure_live_stage_outputs(document: Document) -> None:
    for stage_id, profile in STAGE_PROFILES.items():
        status = StageOutput.Status.BLOCKED if profile.future else StageOutput.Status.NOT_STARTED
        StageOutput.objects.get_or_create(
            document=document,
            stage=stage_id,
            defaults={"display_title": profile.label, "status": status},
        )


def _document_home_payload(document: Document | None) -> dict[str, Any] | None:
    if not document:
        return None
    return {
        "doc_id": document.doc_id,
        "title": document.title,
        "source_file": document.source_file,
        "document_type": document.document_type,
        "normalized_date": document.normalized_date,
        "archival_reference": document.archival_reference,
        "workflow_source": document.metadata.get("workflow_source", ""),
        "placeholder_note": document.metadata.get("placeholder_note", ""),
        "upload_extraction_status": document.metadata.get("upload_extraction_status", ""),
        "upload_extraction_note": document.metadata.get("upload_extraction_note", ""),
        "upload_text_format": document.metadata.get("upload_text_format", ""),
        "clauses_count": document.clauses.count(),
        "created_at": document.created_at,
    }


def _working_source_text_for_document(document: Document) -> str:
    metadata_text = _normalize_working_source_text(str((document.metadata or {}).get("working_source_text", "")))
    if metadata_text:
        return metadata_text
    return _join_working_text_segments([clause.text for clause in document.clauses.all()])


def _replace_document_working_source_text(document: Document, text: str) -> None:
    working_source_text = _normalize_working_source_text(text)
    document.clauses.all().delete()
    Clause.objects.create(
        document=document,
        clause_id="0001",
        sequence=1,
        text=working_source_text,
    )
    metadata = dict(document.metadata or {})
    metadata.update(
        {
            "working_source_text": working_source_text,
            "working_source_text_source": "manual_full_text_edit",
            "working_source_text_updated_at": timezone.now().isoformat(),
            "working_source_paragraph_count": _paragraph_count(working_source_text),
            "working_text_segment_count": 1,
        }
    )
    document.metadata = metadata
    document.save(update_fields=["metadata", "updated_at"])


def _clear_source_text_change_flag(document: Document) -> None:
    metadata = dict(document.metadata or {})
    changed = False
    for key in (
        "source_text_changed_after_workflow_start",
        "source_text_changed_at",
        "source_text_changed_clause",
        "source_text_change_note",
    ):
        if key in metadata:
            metadata.pop(key, None)
            changed = True
    if changed:
        document.metadata = metadata
        document.save(update_fields=["metadata", "updated_at"])


def _mark_source_text_changed(document: Document, clause_id: str, *, workflow_started: bool) -> None:
    metadata = dict(document.metadata or {})
    metadata["source_text_changed_at"] = timezone.now().isoformat()
    metadata["source_text_changed_clause"] = clause_id
    metadata["source_text_change_note"] = (
        "Local working text was edited. Restart the workflow or rerun affected stages "
        "before relying on previous outputs."
    )
    if workflow_started:
        metadata["source_text_changed_after_workflow_start"] = True
    document.metadata = metadata
    document.save(update_fields=["metadata", "updated_at"])


def _active_document(request: HttpRequest) -> Document | None:
    document_id = request.session.get("active_document_pk")
    if not document_id:
        return None
    return Document.objects.filter(pk=document_id).first()


def _document_payload(document: Document) -> dict[str, Any]:
    working_source_text = _working_source_text_for_document(document)
    metadata = {
        "doc_id": document.doc_id,
        "record_id": document.metadata.get("record_id", ""),
        "archival_reference": document.archival_reference,
        "archive_or_library": document.metadata.get("archive_or_library", ""),
        "title": document.title,
        "document_title": document.metadata.get("document_title", document.title),
        "source_file": document.source_file,
        "document_type": document.document_type,
        "originating_body": document.metadata.get("originating_body", ""),
        "plaintiff": document.metadata.get("plaintiff", ""),
        "defendant": document.metadata.get("defendant", ""),
        "normalized_date": document.normalized_date,
        "topic": document.metadata.get("topic", ""),
        "workflow_source": document.metadata.get("workflow_source", ""),
        "placeholder_outputs": document.metadata.get("placeholder_outputs", True),
        "upload_extraction_status": document.metadata.get("upload_extraction_status", ""),
        "upload_extraction_note": document.metadata.get("upload_extraction_note", ""),
        "upload_text_format": document.metadata.get("upload_text_format", ""),
        "source_text_changed_after_workflow_start": document.metadata.get(
            "source_text_changed_after_workflow_start",
            False,
        ),
        "source_text_changed_at": document.metadata.get("source_text_changed_at", ""),
        "source_text_changed_clause": document.metadata.get("source_text_changed_clause", ""),
        "source_text_change_note": document.metadata.get("source_text_change_note", ""),
        "working_source_paragraph_count": document.metadata.get(
            "working_source_paragraph_count",
            _paragraph_count(working_source_text),
        ),
    }
    clauses = [
        {
            "clause_id": clause.clause_id,
            "label": f"Clause {clause.clause_id}",
            "text": clause.text,
        }
        for clause in document.clauses.all()
    ]
    return {"metadata": metadata, "clauses": clauses, "working_source_text": working_source_text}


def _source_reader_context(
    document: dict[str, Any],
    clauses: list[dict[str, Any]],
    stage_outputs: dict[str, StageOutput],
) -> dict[str, Any]:
    clause_parser = stage_outputs.get("clause_parser")
    parser_status = clause_parser.status if clause_parser else ""
    parser_available = parser_status in {
        StageOutput.Status.CHECKING,
        StageOutput.Status.ACCEPTED,
        StageOutput.Status.NEEDS_RERUN,
    }
    metadata = document.get("metadata", {})
    is_uploaded = metadata.get("workflow_source") == "uploaded_document"
    full_text = _normalize_working_source_text(str(document.get("working_source_text") or ""))
    paragraph_count = _paragraph_count(full_text)
    title = (
        metadata.get("document_title")
        or metadata.get("title")
        or metadata.get("doc_id")
        or "Source document"
    )
    metadata_rows: list[dict[str, str]] = []
    if is_uploaded:
        metadata_rows.append({"label": "Source file", "value": str(metadata.get("source_file") or title)})
    else:
        for label, key in (
            ("DocID", "doc_id"),
            ("RecordID", "record_id"),
            ("Archival Reference", "archival_reference"),
        ):
            value = str(metadata.get(key) or "").strip()
            if value:
                metadata_rows.append({"label": label, "value": value})
    for label, key in (
        ("Document Title", "document_title"),
        ("Document Type", "document_type"),
        ("Date", "normalized_date"),
    ):
        value = str(metadata.get(key) or "").strip()
        if value and value != title:
            metadata_rows.append({"label": label, "value": value})
    return {
        "eyebrow": "Source reader",
        "title": title,
        "source_label": "Uploaded source document" if is_uploaded else "Sample source document",
        "full_text": full_text,
        "metadata_rows": metadata_rows,
        "paragraph_count": paragraph_count,
        "segment_count": len(clauses),
        "segment_count_label": "Working source text",
        "reading_label": f"{paragraph_count} paragraph{'s' if paragraph_count != 1 else ''}",
        "parser_available": parser_available,
        "parser_status": parser_status,
        "parser_status_label": clause_parser.get_status_display() if clause_parser else "Not started",
        "parser_note": (
            "Clause Parser output is available as a downstream view for this local workflow."
            if parser_available
            else "Clause Reader appears after the Clause Parser stage is checked or continued."
        ),
        "extraction_note": metadata.get("upload_extraction_note", ""),
        "extraction_status": metadata.get("upload_extraction_status", ""),
        "edit_note": (
            "Edits change the extracted/plain working text used by this prototype. "
            "They do not edit the original DOCX/PDF file or rerun chatbot outputs."
        ),
    }


def _build_clause_html(clause: dict[str, Any], annotations: list[dict[str, Any]]) -> str:
    text = str(clause.get("text") or "")
    matches: list[tuple[int, int, dict[str, Any]]] = []
    lower_text = text.lower()
    for annotation in annotations:
        if str(clause.get("clause_id")) not in [str(item) for item in annotation.get("clause_ids", [])]:
            continue
        phrase = str(annotation.get("evidence_form") or annotation.get("headword") or "").strip()
        if not phrase:
            continue
        start = lower_text.find(phrase.lower())
        if start < 0:
            continue
        matches.append((start, start + len(phrase), annotation))

    matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    filtered: list[tuple[int, int, dict[str, Any]]] = []
    last_end = -1
    for start, end, annotation in matches:
        if start < last_end:
            continue
        filtered.append((start, end, annotation))
        last_end = end

    parts: list[str] = []
    cursor = 0
    for start, end, annotation in filtered:
        parts.append(escape(text[cursor:start]))
        title = (
            f"{annotation.get('type', 'Tag')}: {annotation.get('headword', '')} "
            f"({annotation.get('stable_id', 'provisional')})"
        ).strip()
        detail = json.dumps(
            {
                "type": annotation.get("type"),
                "headword": annotation.get("headword"),
                "stable_id": annotation.get("stable_id"),
                "stage": annotation.get("stage"),
                "review": annotation.get("review_state"),
                "note": annotation.get("note"),
            },
            ensure_ascii=True,
        )
        parts.append(
            '<span class="annotation" tabindex="0" '
            f'data-detail="{escape(detail, quote=True)}" title="{escape(title, quote=True)}">'
            f"{escape(text[start:end])}</span>"
        )
        cursor = end
    parts.append(escape(text[cursor:]))
    return mark_safe("".join(parts))


def _clause_related_items(items: list[dict[str, Any]], clause_id: str) -> list[dict[str, Any]]:
    return [
        item
        for item in items
        if clause_id in [str(value) for value in item.get("clause_ids", [])]
    ]


def _review_summary(proposals: list[NewIdProposal]) -> dict[str, int]:
    return {
        "total": len(proposals),
        "pending": sum(1 for proposal in proposals if proposal.status == NewIdProposal.Status.PENDING),
        "approved": sum(1 for proposal in proposals if proposal.status == NewIdProposal.Status.APPROVED),
        "rejected": sum(1 for proposal in proposals if proposal.status == NewIdProposal.Status.REJECTED),
        "edited": sum(1 for proposal in proposals if proposal.status == NewIdProposal.Status.NEEDS_EDIT),
    }


def _stage_item_review_summary(document: Document, stage_id: str) -> dict[str, int]:
    if stage_id != "entity_registry":
        return {"total": 0, "pending": 0, "approved": 0, "rejected": 0, "edited": 0}
    return _review_summary(list(document.new_id_proposals.all()))


def _approve_pending_item_reviews(stage_output: StageOutput) -> int:
    if stage_output.stage != "entity_registry":
        return 0
    return NewIdProposal.objects.filter(
        document=stage_output.document,
        status=NewIdProposal.Status.PENDING,
    ).update(
        status=NewIdProposal.Status.APPROVED,
        updated_at=timezone.now(),
    )


def _proposal_payload(proposal: NewIdProposal) -> dict[str, Any]:
    return {
        "proposed_id": proposal.proposed_id,
        "type": proposal.record_type,
        "headword": proposal.headword,
        "evidence_form": proposal.evidence_form,
        "source_clause": proposal.source_clause.clause_id if proposal.source_clause else proposal.payload.get("source_clause_label", ""),
        "status": proposal.status,
        "reviewer_note": proposal.reviewer_note,
    }


def _stage_outputs_for_document(document: Document) -> dict[str, StageOutput]:
    return {
        stage_output.stage: stage_output
        for stage_output in document.stage_outputs.all()
    }


def _accepted_stage_ids(stage_outputs: dict[str, StageOutput]) -> set[str]:
    return {
        stage_id
        for stage_id, stage_output in stage_outputs.items()
        if stage_output.status == StageOutput.Status.ACCEPTED
    }


def _missing_accepted_requirements(stage_id: str, stage_outputs: dict[str, StageOutput]) -> list[str]:
    return missing_required_stages(stage_id, _accepted_stage_ids(stage_outputs))


def _first_pending_stage_id(
    ordered_selected_stages: list[str],
    stage_outputs: dict[str, StageOutput],
) -> str:
    for stage_id in ordered_selected_stages:
        stage_output = stage_outputs.get(stage_id)
        if not stage_output or stage_output.status != StageOutput.Status.ACCEPTED:
            return stage_id
    return ""


def _active_workflow_stage_id(
    request: HttpRequest,
    ordered_selected_stages: list[str],
    stage_outputs: dict[str, StageOutput],
) -> str:
    requested_stage = str(request.GET.get("stage") or "").strip()
    all_selected_accepted = bool(ordered_selected_stages) and not _first_pending_stage_id(
        ordered_selected_stages,
        stage_outputs,
    )
    if requested_stage == "final_review" and all_selected_accepted:
        return "final_review"
    if requested_stage in ordered_selected_stages:
        return requested_stage
    pending_stage = _first_pending_stage_id(ordered_selected_stages, stage_outputs)
    if pending_stage:
        return pending_stage
    if all_selected_accepted:
        return "final_review"
    return ""


def _stage_action_detail(stage_output: StageOutput, stage_outputs: dict[str, StageOutput]) -> dict[str, Any]:
    profile = STAGE_PROFILES[stage_output.stage]
    missing_accepted = _missing_accepted_requirements(stage_output.stage, stage_outputs)
    blocked_reason = ""
    if profile.future:
        blocked_reason = "This module is future work and cannot be continued in the current prototype."
    elif missing_accepted:
        blocked_reason = (
            "Continue is available after accepting outputs from: "
            f"{', '.join(stage_labels(missing_accepted))}."
        )
    elif stage_output.status == StageOutput.Status.NEEDS_RERUN:
        blocked_reason = "A chatbot edit request is saved for this stage. Review the request before continuing."
    elif stage_output.status == StageOutput.Status.BLOCKED:
        blocked_reason = (
            "This stage was previously blocked. Continue again to pass it forward "
            "now that prerequisite outputs are accepted."
        )
    return {
        "missing_accepted": missing_accepted,
        "missing_accepted_labels": stage_labels(missing_accepted),
        "can_continue": not profile.future
        and not missing_accepted
        and stage_output.status != StageOutput.Status.NEEDS_RERUN,
        "blocked_reason": blocked_reason,
    }


def _stage_review_notes(stage_output: StageOutput) -> list[ReviewNote]:
    return list(
        stage_output.reviewnote_set.select_related("clause")
        .filter(requested_action="chatbot_rerun_request")
        .order_by("-created_at")[:5]
    )


def _workflow_stage_items(
    *,
    ordered_selected_stages: list[str],
    active_stage_id: str,
    selected_toggles: set[str],
    selected_clause_id: str,
    stage_outputs: dict[str, StageOutput],
) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for index, stage_id in enumerate(ordered_selected_stages, start=1):
        profile = STAGE_PROFILES[stage_id]
        stage_output = stage_outputs.get(stage_id)
        status = stage_output.status if stage_output else StageOutput.Status.NOT_STARTED
        is_active = stage_id == active_stage_id
        is_complete = status == StageOutput.Status.ACCEPTED
        can_open = is_active or is_complete
        state = "active" if is_active else "complete" if is_complete else "queued"
        if status in {StageOutput.Status.BLOCKED, StageOutput.Status.NEEDS_RERUN}:
            state = "warning" if is_active else state
        items.append(
            {
                "index": index,
                "stage_id": stage_id,
                "label": profile.label,
                "input_summary": profile.input_summary,
                "status": status,
                "status_label": stage_output.get_status_display() if stage_output else "Not started",
                "state": state,
                "is_active": is_active,
                "is_complete": is_complete,
                "can_open": can_open,
                "url": _workbench_url(
                    selected_toggles=selected_toggles,
                    selected_clause_id=selected_clause_id,
                    stage_id=stage_id,
                )
                if can_open
                else "",
            }
        )
    return items


def _stage_cards(
    *,
    selected_toggles: set[str],
    selected_stages: set[str],
    active_stage_id: str,
    selected_clause_id: str,
    selected_clause: dict[str, Any],
    clause_annotations: list[dict[str, Any]],
    clause_events: list[dict[str, Any]],
    stage_outputs: dict[str, StageOutput],
    workflow_auto_url: str,
) -> list[dict[str, Any]]:
    cards: list[dict[str, Any]] = []
    for stage_id in STAGE_PROFILES:
        if stage_id != active_stage_id:
            continue
        stage_output = stage_outputs.get(stage_id)
        if not stage_output:
            continue
        profile = STAGE_PROFILES[stage_id]
        payload = stage_output.payload or {}
        action_detail = _stage_action_detail(stage_output, stage_outputs)
        card: dict[str, Any] = {
            "stage_id": stage_id,
            "profile": profile,
            "stage_output": stage_output,
            "status": stage_output.status,
            "status_label": stage_output.get_status_display(),
            "input_summary": profile.input_summary,
            "requires_labels": stage_labels(profile.requires),
            "checker_note": STAGE_CHECKER_NOTES.get(stage_id, ""),
            "action_detail": action_detail,
            "review_notes": _stage_review_notes(stage_output),
            "selected_clause_id": selected_clause_id,
            "show_output": stage_output.status
            in {
                StageOutput.Status.CHECKING,
                StageOutput.Status.ACCEPTED,
                StageOutput.Status.NEEDS_RERUN,
            },
            "workflow_auto_url": workflow_auto_url,
            "runner": payload.get("runner", {}),
            "raw_output": stage_output.raw_output,
        }
        if stage_id == "summary_keywords":
            card["show_summary"] = "summary_keywords" in selected_toggles
            card["show_keywords"] = "summary_keywords" in selected_toggles
            card["summary"] = payload.get("summary", {})
            card["keywords"] = payload.get("keywords", {})
        elif stage_id == "entity_registry":
            card["clause_annotations"] = clause_annotations
            card["entity_registry"] = payload.get("entity_registry", [])
        elif stage_id == "clause_parser":
            card["selected_clause_text"] = selected_clause.get("text", "")
            card["notice"] = payload.get("notice", "")
        elif stage_id == "occurrences_registry":
            card["clause_events"] = clause_events
            card["event_lookup_candidates"] = payload.get("event_lookup_candidates", [])
        elif stage_id == "tag_assembler":
            card["placeholder"] = payload.get("placeholder", "")
            card["full_tagset_placeholder"] = payload.get("full_tagset_placeholder", "")
        elif stage_id == "key_narrative":
            card["future_notice"] = payload.get("future_notice", "")
        cards.append(card)
    return cards


def _event_headword_review_highlighted_clause_html(clause_text: str, event_cut_text: str) -> tuple[Any, bool]:
    if not event_cut_text:
        return mark_safe(escape(clause_text)), False
    start = clause_text.find(event_cut_text)
    if start < 0:
        return mark_safe(escape(clause_text)), False
    end = start + len(event_cut_text)
    html = (
        f"{escape(clause_text[:start])}"
        f'<mark class="eventcut-highlight">{escape(clause_text[start:end])}</mark>'
        f"{escape(clause_text[end:])}"
    )
    return mark_safe(html), True


def _event_headword_review_context_excerpt_html(
    clause_text: str,
    event_cut_text: str,
    *,
    radius: int = 120,
) -> tuple[Any, bool]:
    if not event_cut_text:
        return mark_safe(""), False
    start = clause_text.find(event_cut_text)
    if start < 0:
        return mark_safe(escape(event_cut_text)), False
    excerpt_start = max(0, start - radius)
    excerpt_end = min(len(clause_text), start + len(event_cut_text) + radius)
    excerpt = clause_text[excerpt_start:excerpt_end]
    local_start = start - excerpt_start
    local_end = local_start + len(event_cut_text)
    prefix = "... " if excerpt_start > 0 else ""
    suffix = " ..." if excerpt_end < len(clause_text) else ""
    html = (
        f"{prefix}{escape(excerpt[:local_start])}"
        f'<mark class="eventcut-highlight">{escape(excerpt[local_start:local_end])}</mark>'
        f"{escape(excerpt[local_end:])}{suffix}"
    )
    return mark_safe(html), True


def _event_headword_review_candidate_view(candidate: dict[str, Any]) -> dict[str, Any]:
    snippets = list(candidate.get("matched_snippets") or [])
    snippet_preview = ""
    if snippets:
        snippet_preview = str(snippets[0].get("snippet_text") or "")
    return {
        **candidate,
        "snippet_preview": snippet_preview,
        "detail_id": f"event-headword-candidate-{candidate.get('event_id', '')}-{candidate.get('rank', '')}",
    }


def _event_headword_review_candidate_summary(candidate: dict[str, Any] | None) -> dict[str, Any] | None:
    if not candidate:
        return None
    return {
        "event_id": str(candidate.get("event_id") or ""),
        "headword": str(candidate.get("headword") or ""),
        "rank": candidate.get("rank"),
    }


def _event_headword_review_candidate_by_rank(item: dict[str, Any], rank: str) -> dict[str, Any] | None:
    try:
        selected_rank = int(rank)
    except (TypeError, ValueError):
        return None
    for candidate in item.get("candidate_package", {}).get("candidates") or []:
        if candidate.get("rank") == selected_rank:
            return candidate
    return None


def _event_headword_review_current_proposal(item: dict[str, Any]) -> dict[str, str] | None:
    proposal = (item.get("chooser_result") or {}).get("proposed_new_headword") or {}
    if not proposal.get("suggested") and not proposal.get("headword"):
        decision = item.get("review_decision") or {}
        proposal = decision.get("proposed_headword") or {}
    headword = str(proposal.get("headword") or "").strip()
    if not headword:
        return None
    return {
        "headword": headword,
        "definition_hint": str(proposal.get("definition_hint") or "").strip(),
        "reason": str(proposal.get("reason") or "").strip(),
    }


def _event_headword_review_action_controls(review_state: str, *, has_selected_candidate: bool) -> dict[str, Any]:
    controls: dict[str, Any] = {
        "simple_actions": [],
        "show_candidate_selector": False,
        "candidate_action_label": "Choose candidate",
        "show_proposal_form": False,
        "proposal_summary": "Propose new headword",
    }
    if review_state == "llm_selected_candidate":
        controls["simple_actions"] = [
            {"action": "accept_suggestion", "label": "Accept suggestion", "class": "approve"},
            {"action": "hold_for_later", "label": "Hold for later", "class": "hold"},
        ]
        controls["show_candidate_selector"] = True
        controls["candidate_action_label"] = "Edit / choose another"
        controls["show_proposal_form"] = True
    elif review_state == "needs_historian_review":
        if has_selected_candidate:
            controls["simple_actions"].append(
                {"action": "accept_suggestion", "label": "Accept reviewed candidate", "class": "approve"}
            )
        controls["simple_actions"].append({"action": "hold_for_later", "label": "Hold for later", "class": "hold"})
        controls["show_candidate_selector"] = True
        controls["candidate_action_label"] = "Choose candidate"
        controls["show_proposal_form"] = True
    elif review_state == "llm_none_fit":
        controls["simple_actions"] = [
            {"action": "reject_all_candidates", "label": "Mark no suitable candidate", "class": "reject"},
            {"action": "hold_for_later", "label": "Hold for later", "class": "hold"},
        ]
        controls["show_candidate_selector"] = True
        controls["candidate_action_label"] = "Choose from waiting list"
        controls["show_proposal_form"] = True
    elif review_state == "proposed_new_headword":
        controls["simple_actions"] = [
            {"action": "keep_proposal_pending", "label": "Keep proposal pending review", "class": "proposal"},
            {"action": "hold_for_later", "label": "Hold for later", "class": "hold"},
        ]
        controls["show_candidate_selector"] = True
        controls["candidate_action_label"] = "Choose existing candidate instead"
    return controls


def _event_headword_review_apply_local_overrides(payload: dict[str, Any], overrides: dict[str, Any]) -> dict[str, Any]:
    if not overrides:
        return payload
    payload = deepcopy(payload)
    for item in payload.get("event_review_items", []):
        override = overrides.get(str(item.get("item_id") or ""))
        if not isinstance(override, dict):
            continue
        for key in ("review_state", "review_decision", "audit_log"):
            if key in override:
                item[key] = deepcopy(override[key])
    return payload


def _event_headword_review_apply_review_action(
    payload: dict[str, Any],
    *,
    item_id: str,
    action: str,
    candidate_rank: str = "",
    proposed_headword: str = "",
    definition_hint: str = "",
    reviewer_note: str = "",
    created_at: str = "",
    decision_id: str = "",
) -> tuple[dict[str, Any], dict[str, Any]]:
    payload = deepcopy(payload)
    created_at = created_at or timezone.now().isoformat()
    decision_id = decision_id or f"local-review-{uuid4().hex}"
    reviewer_note = str(reviewer_note or "").strip()
    target_item = None
    for item in payload.get("event_review_items", []):
        if str(item.get("item_id") or "") == str(item_id):
            target_item = item
            break
    if target_item is None:
        raise ValueError("Unknown EventCut review item")

    previous_state = str(target_item.get("review_state") or "")
    chooser = target_item.get("chooser_result") or {}
    selected_candidate = None
    proposed = None
    rejected_candidate_scope = ""
    requires_similarity_check = False
    requires_historian_review = False
    requires_followup = False

    if action == "accept_suggestion":
        if previous_state == "llm_none_fit":
            raise ValueError("No suggested candidate is available for this item")
        selected_candidate = _event_headword_review_candidate_summary(chooser.get("selected_candidate") or {})
        if not selected_candidate or not selected_candidate.get("event_id"):
            raise ValueError("No suggested candidate is available for this item")
        next_state = "accepted_llm_choice"
    elif action == "choose_alternate":
        selected_candidate = _event_headword_review_candidate_summary(_event_headword_review_candidate_by_rank(target_item, candidate_rank))
        if not selected_candidate:
            raise ValueError("Choose a candidate from the waiting list")
        next_state = "accepted_alternate_candidate"
    elif action == "propose_new_headword":
        headword = str(proposed_headword or "").strip()
        if not headword:
            raise ValueError("Proposed headword is required")
        proposed = {
            "headword": headword,
            "definition_hint": str(definition_hint or "").strip(),
            "reason": reviewer_note,
        }
        next_state = "proposed_new_headword_pending_review"
        requires_similarity_check = True
        requires_historian_review = True
    elif action == "keep_proposal_pending":
        proposed = _event_headword_review_current_proposal(target_item)
        if not proposed:
            raise ValueError("No provisional headword is available for this item")
        next_state = "proposed_new_headword_pending_review"
        requires_similarity_check = True
        requires_historian_review = True
    elif action == "reject_all_candidates":
        selected_candidate = None
        next_state = "rejected_all_candidates"
        rejected_candidate_scope = "top20"
        requires_followup = True
    elif action == "hold_for_later":
        selected_candidate = None
        next_state = "held_for_later"
    else:
        raise ValueError("Unknown review action")

    decision: dict[str, Any] = {
        "decision_id": decision_id,
        "action": action,
        "review_state": next_state,
        "selected_candidate": selected_candidate,
        "proposed_headword": proposed,
        "reviewer_note": reviewer_note,
        "actor": "local_user",
        "created_at": created_at,
        "source": "sample_workflow_local_review",
        "not_written_to_official_event_list": True,
    }
    if requires_similarity_check:
        decision["requires_similarity_check"] = True
    if requires_historian_review:
        decision["requires_historian_review"] = True
    if rejected_candidate_scope:
        decision["rejected_candidate_scope"] = rejected_candidate_scope
        decision["requires_followup"] = requires_followup

    audit_event = {
        "timestamp": created_at,
        "actor": "local_user",
        "action": action,
        "previous_state": previous_state,
        "next_state": next_state,
        "item_id": str(target_item.get("item_id") or ""),
        "event_cut_id": str((target_item.get("event_cut") or {}).get("event_cut_id") or ""),
        "selected_event_id": str((selected_candidate or {}).get("event_id") or ""),
        "selected_headword": str((selected_candidate or proposed or {}).get("headword") or ""),
        "selected_rank": (selected_candidate or {}).get("rank"),
        "note": reviewer_note,
    }
    target_item["review_state"] = next_state
    target_item["review_decision"] = decision
    target_item["audit_log"] = list(target_item.get("audit_log") or []) + [audit_event]
    return payload, target_item


def _event_headword_review_item_view(
    *,
    item: dict[str, Any],
    selected_toggles: set[str],
    is_selected: bool,
) -> dict[str, Any]:
    event_cut = item.get("event_cut", {})
    clause = item.get("clause", {})
    chooser = item.get("chooser_result", {})
    decision = str(chooser.get("decision") or "")
    review_state = str(item.get("review_state") or "")
    clause_id = str(clause.get("clause_id") or "")
    event_cut_text = str(event_cut.get("text_verbatim") or "")
    clause_text = str(clause.get("text_verbatim") or "")
    highlighted_clause_html, highlight_found = _event_headword_review_highlighted_clause_html(clause_text, event_cut_text)
    context_excerpt_html, context_excerpt_found = _event_headword_review_context_excerpt_html(clause_text, event_cut_text)
    url = _workbench_url(
        selected_toggles=selected_toggles,
        selected_clause_id=clause_id,
        stage_id="occurrences_registry",
    )
    url = f"{url}&{urlencode({'event_headword_review_item': str(item.get('item_id') or '')})}"
    selected_candidate = chooser.get("selected_candidate") or {}
    proposed_headword = chooser.get("proposed_new_headword") or {}
    candidates = [
        _event_headword_review_candidate_view(candidate)
        for candidate in (item.get("candidate_package", {}).get("candidates") or [])
    ]
    selected_candidate_view = next(
        (candidate for candidate in candidates if candidate.get("is_selected_by_mock_chooser")),
        None,
    )
    headword_preview = ""
    if selected_candidate.get("headword"):
        headword_preview = str(selected_candidate.get("headword") or "")
    elif proposed_headword.get("headword"):
        headword_preview = str(proposed_headword.get("headword") or "")
    review_decision = item.get("review_decision") or {}
    recorded_label = EVENT_HEADWORD_REVIEW_RECORDED_STATE_LABELS.get(review_state, "")
    decision_label = (
        recorded_label
        or EVENT_HEADWORD_REVIEW_DECISION_LABELS.get(decision, decision.replace("_", " ").title())
    )
    decision_class = EVENT_HEADWORD_REVIEW_DECISION_CLASSES.get(review_state) or EVENT_HEADWORD_REVIEW_DECISION_CLASSES.get(decision, "review")
    state_class = EVENT_HEADWORD_REVIEW_STATE_CLASSES.get(review_state, "review")
    action_controls = _event_headword_review_action_controls(
        review_state,
        has_selected_candidate=bool(selected_candidate.get("event_id")),
    )
    return {
        **item,
        "is_selected": is_selected,
        "url": url,
        "clause_id": clause_id,
        "event_cut_text": event_cut_text,
        "highlighted_clause_html": highlighted_clause_html,
        "highlight_found": highlight_found,
        "context_excerpt_html": context_excerpt_html,
        "context_excerpt_found": context_excerpt_found,
        "decision": decision,
        "decision_label": decision_label,
        "decision_class": decision_class,
        "state_label": EVENT_HEADWORD_REVIEW_STATE_LABELS.get(review_state, review_state.replace("_", " ").title()),
        "state_class": state_class,
        "confidence_label": str(chooser.get("confidence_label") or "").replace("_", " ").title(),
        "reason": chooser.get("reason", ""),
        "review_decision": review_decision,
        "has_recorded_decision": bool(review_decision),
        "recorded_decision_label": recorded_label,
        "selected_candidate": selected_candidate,
        "selected_candidate_view": selected_candidate_view,
        "selected_candidate_label": (
            f"{selected_candidate.get('event_id')} | {selected_candidate.get('headword')}"
            if selected_candidate.get("event_id")
            else ""
        ),
        "headword_preview": headword_preview,
        "candidate_assessment": chooser.get("candidate_assessment") or [],
        "none_fit": chooser.get("none_fit") or {},
        "needs_historian_review": chooser.get("needs_historian_review") or {},
        "proposed_new_headword": proposed_headword,
        "proposed_headword_label": str(proposed_headword.get("headword") or ""),
        "candidates": candidates,
        "remaining_candidates": candidates[5:],
        "action_controls": action_controls,
    }


def _event_headword_review_context(
    *,
    document: Document,
    stage_outputs: dict[str, StageOutput],
    selected_toggles: set[str],
    selected_item_id: str,
    active_stage_id: str,
    local_overrides: dict[str, Any] | None = None,
    notice: str = "",
) -> dict[str, Any]:
    if not _is_sample_document(document):
        return {"available": False, "reason": "Event Headword Review is available only for the sample guided demo."}
    if active_stage_id != "occurrences_registry":
        return {"available": False, "reason": "Event Headword Review appears in the Occurrences Registry stage."}

    stage_output = stage_outputs.get("occurrences_registry")
    stage_payload = stage_output.payload if stage_output else {}
    payload = (stage_payload or {}).get("event_headword_review") or _load_event_headword_review_fixture()
    payload = _event_headword_review_apply_local_overrides(payload, local_overrides or {})
    raw_items = list(payload.get("event_review_items") or [])
    if not raw_items:
        return {"available": False, "reason": "No Event headword review fixture items are available."}

    if selected_item_id not in {str(item.get("item_id") or "") for item in raw_items}:
        selected_item_id = str(raw_items[0].get("item_id") or "")

    items = [
        _event_headword_review_item_view(
            item=item,
            selected_toggles=selected_toggles,
            is_selected=str(item.get("item_id") or "") == selected_item_id,
        )
        for item in raw_items
    ]
    selected_item = next(item for item in items if item["is_selected"])
    bulk_accept_remaining_count = sum(
        1 for item in raw_items if str(item.get("review_state") or "") == "llm_selected_candidate"
    )
    return {
        "available": True,
        "payload": payload,
        "items": items,
        "selected_item": selected_item,
        "selected_item_id": selected_item_id,
        "counts": payload.get("counts", {}),
        "baseline": payload.get("baseline", {}),
        "provenance": payload.get("provenance", {}),
        "notice": notice,
        "stage_status": stage_output.status if stage_output else "",
        "stage_status_label": stage_output.get_status_display() if stage_output else "",
        "bulk_accept_remaining_count": bulk_accept_remaining_count,
    }


def _event_headword_review_redirect_url(
    *,
    selected_toggles: set[str],
    selected_clause_id: str,
    item_id: str,
) -> str:
    url = _workbench_url(
        selected_toggles=selected_toggles,
        selected_clause_id=selected_clause_id,
        stage_id="occurrences_registry",
    )
    return f"{url}&{urlencode({'event_headword_review_item': item_id})}"


def _store_event_headword_review_session_override(
    request: HttpRequest,
    item: dict[str, Any],
    *,
    notice: str,
) -> None:
    item_id = str(item.get("item_id") or "")
    overrides = dict(request.session.get(EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY, {}))
    overrides[item_id] = {
        "review_state": item.get("review_state"),
        "review_decision": item.get("review_decision"),
        "audit_log": item.get("audit_log") or [],
    }
    request.session[EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY] = overrides
    request.session[EVENT_HEADWORD_REVIEW_SESSION_NOTICE_KEY] = notice


def _clear_event_headword_review_session_override(request: HttpRequest, item_id: str) -> None:
    overrides = dict(request.session.get(EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY, {}))
    if item_id in overrides:
        overrides.pop(item_id, None)
        request.session[EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY] = overrides


def _handle_event_headword_review_action(request: HttpRequest, document: Document) -> HttpResponse:
    if not _is_sample_document(document):
        raise Http404("Event Headword Review actions are available only for the sample guided demo.")

    outputs = _load_json_fixture("sample_outputs.json")
    try:
        _seed_stage_outputs_for_document(document, outputs)
        stage_output = StageOutput.objects.get(document=document, stage="occurrences_registry")
        stage_payload = dict(stage_output.payload or {})
        review_payload = stage_payload.get("event_headword_review") or _load_event_headword_review_fixture()
    except DatabaseError:
        stage_output = None
        stage_payload = {"occurrences_registry": outputs.get("events", [])}
        review_payload = _event_headword_review_apply_local_overrides(
            _load_event_headword_review_fixture(),
            dict(request.session.get(EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY, {})),
        )

    item_id = str(request.POST.get("event_headword_review_item") or "").strip()
    action_name = str(request.POST.get("event_headword_review_action") or "").strip()
    selected_clause_id = str(request.POST.get("target_clause") or request.POST.get("clause") or "").strip()
    try:
        if action_name == "accept_remaining":
            updated_review_payload = deepcopy(review_payload)
            updated_item = None
            accepted_count = 0
            for candidate_item in list(updated_review_payload.get("event_review_items") or []):
                if str(candidate_item.get("review_state") or "") != "llm_selected_candidate":
                    continue
                updated_review_payload, updated_item = _event_headword_review_apply_review_action(
                    updated_review_payload,
                    item_id=str(candidate_item.get("item_id") or ""),
                    action="accept_suggestion",
                    reviewer_note="Accepted by bulk accept-remaining action.",
                )
                accepted_count += 1
            if updated_item is None:
                raise ValueError("No untouched suggested Event headwords remain")
        else:
            updated_review_payload, updated_item = _event_headword_review_apply_review_action(
                review_payload,
                item_id=item_id,
                action=action_name,
                candidate_rank=str(request.POST.get("candidate_rank") or ""),
                proposed_headword=str(request.POST.get("proposed_headword") or ""),
                definition_hint=str(request.POST.get("definition_hint") or ""),
                reviewer_note=str(request.POST.get("reviewer_note") or ""),
            )
    except ValueError as exc:
        request.session[EVENT_HEADWORD_REVIEW_SESSION_NOTICE_KEY] = str(exc)
        return redirect(_safe_next_url(request))

    stage_payload["event_headword_review"] = updated_review_payload
    if stage_output is not None:
        try:
            stage_output.payload = stage_payload
            provenance = dict(stage_output.provenance or {})
            provenance["event_headword_review_local_review_updated_at"] = timezone.now().isoformat()
            provenance["event_headword_review_local_review_last_action"] = action_name
            provenance["real_chatbot_execution"] = False
            provenance["postgresql_commit"] = False
            stage_output.provenance = provenance
            stage_output.save(update_fields=["payload", "provenance", "updated_at"])
            _clear_event_headword_review_session_override(request, item_id)
            request.session[EVENT_HEADWORD_REVIEW_SESSION_NOTICE_KEY] = (
                f"Accepted {accepted_count} remaining Event headword suggestions; existing edits and other decisions were preserved."
                if action_name == "accept_remaining"
                else "Local review decision saved in the sample workflow."
            )
        except DatabaseError:
            _store_event_headword_review_session_override(
                request,
                updated_item,
                notice=(
                    "Database write failed, so this local review decision is shown from a session fallback. "
                    "Fix SQLite write permissions to persist it in StageOutput."
                ),
            )
    else:
        _store_event_headword_review_session_override(
            request,
            updated_item,
            notice=(
                "Database write failed, so this local review decision is shown from a session fallback. "
                "Fix SQLite write permissions to persist it in StageOutput."
            ),
        )

    if not selected_clause_id:
        selected_clause_id = str((updated_item.get("clause") or {}).get("clause_id") or "")
    selected_toggles = _selected_toggles(request)
    return redirect(
        _event_headword_review_redirect_url(
            selected_toggles=selected_toggles,
            selected_clause_id=selected_clause_id,
            item_id=item_id,
        )
    )


def _stage_output_payload(stage_output: StageOutput) -> dict[str, Any]:
    return {
        "stage_id": stage_output.stage,
        "stage": stage_output.stage,
        "label": stage_output.display_title,
        "status": stage_output.status,
        "raw_output": stage_output.raw_output,
        "payload": stage_output.payload,
        "provenance": stage_output.provenance,
        "updated_at": stage_output.updated_at.isoformat() if stage_output.updated_at else "",
    }


def _stage_note_payload(note: ReviewNote) -> dict[str, Any]:
    return {
        "stage": note.stage_output.stage if note.stage_output else "",
        "stage_label": note.stage_output.display_title if note.stage_output else "",
        "clause": note.clause.clause_id if note.clause else "",
        "note": note.note,
        "requested_action": note.requested_action,
        "created_at": note.created_at.isoformat() if note.created_at else "",
    }


def home(request: HttpRequest) -> HttpResponse:
    sample_document = _load_json_fixture("sample_document.json")
    sample_metadata = sample_document.get("metadata", {})

    if request.method == "POST":
        action = str(request.POST.get("action") or "").strip()
        if action == "event_headword_review_action":
            active_document = _active_document(request)
            if not active_document or not _workflow_started(request):
                return redirect(reverse("tagger:home"))
            return _handle_event_headword_review_action(request, active_document)
        if action == "use_sample":
            document = _ensure_sample_document(reset_review_state=True)
            request.session["active_document_pk"] = document.pk
            request.session["workflow_started"] = False
            _set_home_notice(request, "Sample document is ready. Start the tagging workflow when you are ready.")
            return redirect(reverse("tagger:home"))
        if action == "upload_document":
            uploaded_file = request.FILES.get("document_file")
            if uploaded_file:
                document = _create_uploaded_document(uploaded_file)
                request.session["active_document_pk"] = document.pk
                request.session["workflow_started"] = False
                _set_home_notice(request, f"Uploaded document ready: {document.source_file}")
            else:
                _set_home_notice(request, "Choose a document file before opening it.")
            return redirect(reverse("tagger:home"))
        if action == "start_workflow":
            active_document = _active_document(request)
            if not active_document:
                _set_home_notice(request, "Choose or upload a document before starting the workflow.")
                return redirect(reverse("tagger:home"))
            is_live = (active_document.metadata or {}).get("workflow_source") == "uploaded_document"
            if request.POST.get("reset_workflow") == "1" and not is_live:
                _reset_stage_outputs_for_document(active_document)
                _clear_source_text_change_flag(active_document)
            request.session["workflow_started"] = True
            return redirect(reverse("tagger:live_workbench") if is_live else reverse("tagger:workbench"))
        if action == "clear_document":
            _clear_active_workflow(request)
            _set_home_notice(request, "Document selection cleared.")
            return redirect(reverse("tagger:home"))

    active_document = _active_document(request)
    if active_document and _is_sample_document(active_document):
        _seed_stage_outputs_for_document(active_document, _load_json_fixture("sample_outputs.json"))
    elif active_document:
        _ensure_live_stage_outputs(active_document)
    stage_outputs = list(active_document.stage_outputs.all()) if active_document else []
    progress_count = sum(1 for stage_output in stage_outputs if stage_output.status == StageOutput.Status.ACCEPTED)
    has_progress = progress_count > 0 or any(
        stage_output.status in {StageOutput.Status.CHECKING, StageOutput.Status.NEEDS_RERUN}
        for stage_output in stage_outputs
    )
    context = {
        "nav_active": "home",
        "page_title": "Document Workspace",
        "sample_doc_id": SAMPLE_DOC_ID,
        "sample_metadata": sample_metadata,
        "active_document": active_document,
        "active_document_payload": _document_home_payload(active_document),
        "workflow_started": _workflow_started(request),
        "has_progress": has_progress,
        "progress_count": progress_count,
        "home_notice": request.session.pop("home_notice", ""),
    }
    return render(request, "tagger/home.html", context)


def workbench(request: HttpRequest) -> HttpResponse:
    active_for_route = _active_document(request)
    if (
        request.method == "GET"
        and active_for_route
        and (active_for_route.metadata or {}).get("workflow_source") == "uploaded_document"
    ):
        return redirect(reverse("tagger:live_workbench"))
    outputs = _load_json_fixture("sample_outputs.json")

    if request.method == "POST":
        action = str(request.POST.get("action") or "").strip()
        if action == "use_sample":
            document = _ensure_sample_document(reset_review_state=True)
            request.session["active_document_pk"] = document.pk
            request.session["workflow_started"] = False
            return redirect(reverse("tagger:home"))
        if action == "upload_document":
            uploaded_file = request.FILES.get("document_file")
            if uploaded_file:
                document = _create_uploaded_document(uploaded_file)
                request.session["active_document_pk"] = document.pk
                request.session["workflow_started"] = False
            return redirect(reverse("tagger:home"))
        if action == "clear_document":
            _clear_active_workflow(request)
            return redirect(reverse("tagger:home"))
        if action == "reset_workflow":
            active_document = _active_document(request)
            if active_document:
                _reset_stage_outputs_for_document(active_document)
                _clear_source_text_change_flag(active_document)
            return redirect(reverse("tagger:workbench"))
        if action == "update_source_text":
            active_document = _active_document(request)
            source_text = _normalize_working_source_text(str(request.POST.get("source_text") or ""))
            if active_document and source_text:
                _replace_document_working_source_text(active_document, source_text[:200000])
                _mark_source_text_changed(
                    active_document,
                    "full_document",
                    workflow_started=_workflow_started(request),
                )
            return redirect(_safe_next_url(request))
        if action == "correction_note":
            note_text = str(request.POST.get("correction_note") or "").strip()
            if note_text:
                notes = list(request.session.get("review_notes", []))
                notes.insert(
                    0,
                    {
                        "target_clause": str(request.POST.get("target_clause") or ""),
                        "stage_id": str(request.POST.get("stage_id") or ""),
                        "note": note_text,
                        "created_at": timezone.localtime().strftime("%Y-%m-%d %H:%M"),
                        "status": "Saved as a local review note",
                    },
                )
                request.session["review_notes"] = notes[:8]
            return redirect(_safe_next_url(request))

    active_document = _active_document(request)
    selected_toggles = _selected_toggles(request)
    requested_stages = selected_stage_ids(selected_toggles)
    selected_stages = dependency_closure(requested_stages)
    ordered_selected_stages = ordered_stage_ids(selected_stages)
    toggle_query = _toggle_query(selected_toggles)

    if not active_document or not _workflow_started(request):
        return redirect(reverse("tagger:home"))

    _seed_stage_outputs_for_document(active_document, outputs)
    stage_outputs = _stage_outputs_for_document(active_document)
    document = _document_payload(active_document)
    clauses = list(document.get("clauses", [])) or [{"clause_id": "0001", "text": "No clause text available yet."}]
    selected_clause_id = request.GET.get("clause") or clauses[0].get("clause_id")
    if selected_clause_id not in {str(clause.get("clause_id")) for clause in clauses}:
        selected_clause_id = str(clauses[0].get("clause_id"))

    annotations = list(outputs.get("annotations", []))
    rendered_clauses = []
    for clause in clauses:
        rendered = dict(clause)
        rendered["html"] = _build_clause_html(clause, annotations)
        rendered["is_selected"] = str(clause.get("clause_id")) == str(selected_clause_id)
        rendered_clauses.append(rendered)

    selected_clause = next(
        clause for clause in rendered_clauses if str(clause.get("clause_id")) == str(selected_clause_id)
    )
    proposals = list(active_document.new_id_proposals.select_related("source_clause").all())
    review_summary = _review_summary(proposals)
    warnings = dependency_warnings(requested_stages)
    clause_annotations = _clause_related_items(annotations, selected_clause_id)
    clause_events = _clause_related_items(list(outputs.get("events", [])), selected_clause_id)
    stage_statuses = {stage_id: output.status for stage_id, output in stage_outputs.items()}
    active_stage_id = _active_workflow_stage_id(request, ordered_selected_stages, stage_outputs)
    orchestration_plan = build_orchestration_plan(
        document=active_document,
        requested_stages=requested_stages,
        stage_outputs=stage_outputs,
    )
    source_reader = _source_reader_context(document, rendered_clauses, stage_outputs)
    event_headword_review = _event_headword_review_context(
        document=active_document,
        stage_outputs=stage_outputs,
        selected_toggles=selected_toggles,
        selected_item_id=str(request.GET.get("event_headword_review_item") or ""),
        active_stage_id=active_stage_id,
        local_overrides=dict(request.session.get(EVENT_HEADWORD_REVIEW_SESSION_OVERRIDE_KEY, {})),
        notice=str(request.session.pop(EVENT_HEADWORD_REVIEW_SESSION_NOTICE_KEY, "") or ""),
    )
    final_review_active = active_stage_id == "final_review"
    workflow_auto_url = _workbench_url(
        selected_toggles=selected_toggles,
        selected_clause_id=selected_clause_id,
    )
    workflow_stage_items = _workflow_stage_items(
        ordered_selected_stages=ordered_selected_stages,
        active_stage_id=active_stage_id,
        selected_toggles=selected_toggles,
        selected_clause_id=selected_clause_id,
        stage_outputs=stage_outputs,
    )
    active_stage_label = (
        "Final review"
        if final_review_active
        else STAGE_PROFILES[active_stage_id].label
        if active_stage_id in STAGE_PROFILES
        else "Workflow"
    )
    active_stage_output = stage_outputs.get(active_stage_id)
    active_stage_review_summary = _stage_item_review_summary(active_document, active_stage_id)
    show_new_id_review = final_review_active or (
        active_stage_id == "entity_registry"
        and active_stage_output is not None
        and active_stage_output.status
        in {
            StageOutput.Status.CHECKING,
            StageOutput.Status.ACCEPTED,
            StageOutput.Status.NEEDS_RERUN,
        }
    )
    show_downloads = final_review_active
    source_text_changed = bool(document.get("metadata", {}).get("source_text_changed_after_workflow_start"))

    context = {
        "nav_active": "tagger",
        "page_title": "Tagger Workbench",
        "entry_only": False,
        "document": document,
        "active_document": active_document,
        "outputs": outputs,
        "clauses": rendered_clauses,
        "selected_clause": selected_clause,
        "selected_clause_id": selected_clause_id,
        "source_reader": source_reader,
        "event_headword_review": event_headword_review,
        "selected_toggles": selected_toggles,
        "toggle_query": toggle_query,
        "tagging_controls": control_context(selected_toggles),
        "selected_stages": selected_stages,
        "requested_stages": requested_stages,
        "orchestration_plan": orchestration_plan.as_dict(),
        "stage_profiles": STAGE_PROFILES,
        "stage_options": list(STAGE_PROFILES.values()),
        "dependency_warnings": warnings,
        "pipeline": pipeline_context(
            selected_stages=selected_stages,
            has_document=True,
            review_summary=review_summary,
            stage_statuses=stage_statuses,
            active_stage_id=active_stage_id,
        ),
        "clause_annotations": clause_annotations,
        "clause_events": clause_events,
        "stage_outputs": stage_outputs,
        "stage_cards": _stage_cards(
            selected_toggles=selected_toggles,
            selected_stages=selected_stages,
            active_stage_id=active_stage_id,
            selected_clause_id=selected_clause_id,
            selected_clause=selected_clause,
            clause_annotations=clause_annotations,
            clause_events=clause_events,
            stage_outputs=stage_outputs,
            workflow_auto_url=workflow_auto_url,
        ),
        "stage_statuses": stage_statuses,
        "workflow_stage_items": workflow_stage_items,
        "ordered_selected_stages": ordered_selected_stages,
        "active_stage_id": active_stage_id,
        "active_stage_label": active_stage_label,
        "active_stage_output": active_stage_output,
        "active_stage_review_summary": active_stage_review_summary,
        "final_review_active": final_review_active,
        "workflow_auto_url": workflow_auto_url,
        "source_text_changed": source_text_changed,
        "source_text_change_note": document.get("metadata", {}).get("source_text_change_note", ""),
        "show_new_id_review": show_new_id_review,
        "show_downloads": show_downloads,
        "review_notes": request.session.get("review_notes", []),
        "new_id_proposals": proposals,
        "review_summary": review_summary,
        "edit_proposal_id": int(request.GET.get("edit_id") or 0),
        "proposal_status_choices": NewIdProposal.Status.choices,
        "entity_record_types": sorted(ENTITY_RECORD_TYPES),
        "entity_proposal_notice": request.session.pop("entity_proposal_notice", ""),
        "download_artifacts": [
            ("summary", "Summary JSON"),
            ("entity-registry", "Entity Registry JSON"),
            ("key-events", "Occurrences JSON"),
            ("new-ids", "NEW-ID CSV"),
            ("final-review", "Final Review JSON"),
        ],
    }
    return render(request, "tagger/workbench.html", context)


def update_new_id_proposal(request: HttpRequest, proposal_id: int, action: str) -> HttpResponse:
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])
    proposal = get_object_or_404(NewIdProposal, pk=proposal_id)
    if action == "approve":
        proposal.status = NewIdProposal.Status.APPROVED
    elif action == "reject":
        proposal.status = NewIdProposal.Status.REJECTED
    elif action == "edit":
        proposal.record_type = str(request.POST.get("record_type") or proposal.record_type).strip()
        proposal.headword = str(request.POST.get("headword") or proposal.headword).strip()
        proposal.evidence_form = str(request.POST.get("evidence_form") or proposal.evidence_form).strip()
        proposal.reviewer_note = str(request.POST.get("reviewer_note") or "").strip()
        requested_status = str(request.POST.get("status") or proposal.status)
        valid_statuses = {choice[0] for choice in NewIdProposal.Status.choices}
        proposal.status = requested_status if requested_status in valid_statuses else NewIdProposal.Status.NEEDS_EDIT
    else:
        raise Http404("Unknown review action")
    proposal.save(update_fields=["record_type", "headword", "evidence_form", "reviewer_note", "status", "updated_at"])
    return redirect(_safe_next_url(request))


def create_entity_proposal(request: HttpRequest, stage_output_id: int) -> HttpResponse:
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])
    stage_output = get_object_or_404(
        StageOutput.objects.select_related("document"),
        pk=stage_output_id,
        stage=StageOutput.Stage.ENTITY_REGISTRY,
    )
    clause_id = str(request.POST.get("source_clause") or "").strip()
    clause = Clause.objects.filter(document=stage_output.document, clause_id=clause_id).first()
    try:
        propose_entity(
            stage_output=stage_output,
            record_type=str(request.POST.get("record_type") or ""),
            headword=str(request.POST.get("headword") or ""),
            evidence_form=str(request.POST.get("evidence_form") or ""),
            reviewer_note=str(request.POST.get("reviewer_note") or ""),
            source_clause=clause,
        )
    except ValueError as exc:
        request.session["entity_proposal_notice"] = str(exc)
    else:
        request.session["entity_proposal_notice"] = "Entity proposal added to the review queue."
    return redirect(_safe_next_url(request))


def update_stage_output(request: HttpRequest, stage_output_id: int, action: str) -> HttpResponse:
    if request.method != "POST":
        return HttpResponseNotAllowed(["POST"])
    stage_output = get_object_or_404(
        StageOutput.objects.select_related("document"),
        pk=stage_output_id,
    )
    profile = STAGE_PROFILES.get(stage_output.stage)
    if not profile:
        raise Http404("Unknown stage")

    stage_outputs = _stage_outputs_for_document(stage_output.document)
    stage_outputs[stage_output.stage] = stage_output
    provenance = dict(stage_output.provenance or {})

    if stage_output.status == StageOutput.Status.ACCEPTED and action in {"check", "request-rerun", "run-local", "run-backend-stub"}:
        provenance["read_only_action_ignored_at"] = timezone.now().isoformat()
        stage_output.provenance = provenance
        stage_output.save(update_fields=["provenance", "updated_at"])
        return redirect(_safe_next_url(request))

    if action == "refresh-event-candidates":
        if stage_output.stage != "occurrences_registry":
            raise Http404("Event candidate lookup is available only for Occurrences Registry.")
        source_body = _working_source_text_for_document(stage_output.document)
        payload = dict(stage_output.payload or {})
        payload["event_lookup_candidates"] = build_event_candidate_packages(
            stage_outputs,
            source_body,
        )
        payload["runner"] = {
            "status": "event_candidates_ready",
            "provider": "deterministic_lexical_authority",
            "model": "",
            "message": (
                "Development candidates generated from Events_List_VectorLLM_v1.xlsx. "
                "This is a model-free deterministic lexical contract backend, not dense vector search. "
                "TF-IDF is forbidden."
            ),
        }
        stage_output.payload = payload
        provenance["event_lookup"] = {
            "source": "Events_List_VectorLLM_v1.xlsx",
            "representation": "Headword + Definition + Vector Example",
            "backend": "deterministic_lexical_authority",
            "query_scope": "whole_clause_development_fallback",
            "generated_at": timezone.now().isoformat(),
            "production_vector_search": False,
        }
        stage_output.provenance = provenance
        stage_output.save(update_fields=["payload", "provenance", "updated_at"])
        return redirect(_safe_next_url(request))
    if action in {"run-local", "run-fixture", "run-backend-stub"}:
        missing = _missing_accepted_requirements(stage_output.stage, stage_outputs)
        if missing:
            provenance["last_runner_attempt"] = {
                "status": "blocked",
                "error": (
                    "Accept required outputs first: "
                    f"{', '.join(stage_labels(missing))}."
                ),
                "attempted_at": timezone.now().isoformat(),
            }
            stage_output.provenance = provenance
            stage_output.save(update_fields=["provenance", "updated_at"])
            return redirect(_safe_next_url(request))

        fixture_outputs = _load_json_fixture("sample_outputs.json")
        fixture_payload, fixture_raw = _stage_fixture_payload(
            stage_output.stage,
            stage_output.document,
            fixture_outputs,
        )
        correction = ""
        if stage_output.status == StageOutput.Status.NEEDS_RERUN:
            latest_note = (
                ReviewNote.objects.filter(
                    document=stage_output.document,
                    stage_output=stage_output,
                    requested_action="chatbot_rerun_request",
                )
                .order_by("-created_at")
                .first()
            )
            correction = latest_note.note if latest_note else ""
        result = ChatbotStageRunner().run(
            stage_id=stage_output.stage,
            document=stage_output.document,
            stage_outputs=stage_outputs,
            correction=correction,
            provider={
                "run-fixture": ProviderLabel.FIXTURE.value,
                "run-backend-stub": ProviderLabel.BACKEND_STUB.value,
            }.get(action, ProviderLabel.LOCAL_CPU.value),
            fixture_payload=fixture_payload,
            fixture_raw_output=fixture_raw,
        )
        if (
            stage_output.stage == StageOutput.Stage.ENTITY_REGISTRY
            and result.provider != ProviderLabel.FIXTURE.value
        ):
            persist_entity_registry_attempt(stage_output=stage_output, result=result)
            return redirect(_safe_next_url(request))
        if result.status in {"completed", "fixture"}:
            stage_output.payload = {
                **result.payload,
                "runner": {
                    "status": result.status,
                    "provider": result.provenance.get("provider", ""),
                    "model": result.provenance.get("model", ""),
                "message": (
                    "Fixture fallback restored."
                    if result.provider == ProviderLabel.FIXTURE.value
                    else "Stage generated by the local runner."
                ),
                },
            }
            stage_output.raw_output = result.raw_output
            stage_output.provenance = result.provenance
            stage_output.status = StageOutput.Status.CHECKING
            if (
                stage_output.stage == "clause_parser"
                and result.status == "completed"
                and (result.payload.get("coverage_validation") or {}).get("valid")
            ):
                replace_document_clauses(
                    stage_output.document,
                    list(result.payload.get("clauses") or []),
                )
                stage_output.provenance["clauses_applied_to_document"] = True
        else:
            payload = dict(stage_output.payload or {})
            payload["runner"] = {
                "status": result.status,
                "provider": result.provenance.get("provider", ""),
                "model": result.provenance.get("model", ""),
                "message": result.error or "The stage runner did not complete.",
                "fixture_fallback_preserved": True,
            }
            stage_output.payload = payload
            provenance.update(result.provenance)
            provenance["last_runner_attempt"] = {
                "status": result.status,
                "error": result.error,
                "attempted_at": timezone.now().isoformat(),
                "fixture_fallback_preserved": True,
            }
            stage_output.provenance = provenance
        stage_output.save(
            update_fields=["status", "payload", "raw_output", "provenance", "updated_at"]
        )
        return redirect(_safe_next_url(request))
    if action == "check":
        stage_output.status = StageOutput.Status.CHECKING
        provenance["checked_at"] = timezone.now().isoformat()
        provenance.pop("blocked_reason", None)
    elif action == "approve-pending":
        approved_count = _approve_pending_item_reviews(stage_output)
        provenance["approved_pending_item_reviews_at"] = timezone.now().isoformat()
        provenance["approved_pending_item_reviews_count"] = approved_count
        provenance["approved_pending_item_reviews_scope"] = stage_output.stage
    elif action == "continue":
        missing = _missing_accepted_requirements(stage_output.stage, stage_outputs)
        if profile.future:
            stage_output.status = StageOutput.Status.BLOCKED
            provenance["blocked_reason"] = (
                "This module is future work and cannot be continued in the current prototype."
            )
        elif missing:
            stage_output.status = StageOutput.Status.BLOCKED
            provenance["blocked_reason"] = (
                "Continue blocked until accepted outputs exist for: "
                f"{', '.join(stage_labels(missing))}."
            )
        elif stage_output.status == StageOutput.Status.NEEDS_RERUN:
            provenance["blocked_reason"] = (
                "A chatbot edit request is saved for this stage. Review the request before continuing."
            )
        else:
            approved_count = _approve_pending_item_reviews(stage_output)
            stage_output.status = StageOutput.Status.ACCEPTED
            if stage_output.stage == StageOutput.Stage.ENTITY_REGISTRY:
                payload = dict(stage_output.payload or {})
                entity_review = dict(payload.get("entity_review") or {})
                entity_review["state"] = "approved"
                entity_review["approved_for_downstream"] = True
                payload["entity_review"] = entity_review
                stage_output.payload = payload
                entity_provenance = dict(provenance.get("entity_registry") or {})
                entity_provenance["approved_for_downstream"] = True
                provenance["entity_registry"] = entity_provenance
            provenance["continued_at"] = timezone.now().isoformat()
            provenance["passed_forward_locally"] = True
            provenance["pending_item_reviews_approved_on_continue"] = approved_count
            provenance["continue_acceptance_note"] = (
                "Pending item-level rows for this stage were treated as approved locally; "
                "explicit rejects and edited rows were preserved."
            )
            provenance.pop("blocked_reason", None)
    elif action == "request-rerun":
        note_text = str(request.POST.get("rerun_request") or "").strip()
        target_clause = str(request.POST.get("target_clause") or "").strip()
        clause = Clause.objects.filter(
            document=stage_output.document,
            clause_id=target_clause,
        ).first()
        if note_text:
            ReviewNote.objects.create(
                document=stage_output.document,
                clause=clause,
                stage_output=stage_output,
                note=note_text,
                requested_action="chatbot_rerun_request",
                created_by_label="local prototype reviewer",
            )
        stage_output.status = StageOutput.Status.NEEDS_RERUN
        provenance["last_rerun_request_at"] = timezone.now().isoformat()
        provenance["real_chatbot_rerun"] = False
        provenance.pop("blocked_reason", None)
    else:
        raise Http404("Unknown stage action")

    stage_output.provenance = provenance
    stage_output.save(update_fields=["status", "provenance", "updated_at"])
    return redirect(_safe_next_url(request))


def _active_or_sample_document(request: HttpRequest) -> Document:
    return _active_document(request) or _ensure_sample_document()


def _stage_payload_or_empty(stage_outputs: dict[str, StageOutput], stage_id: str) -> dict[str, Any]:
    stage_output = stage_outputs.get(stage_id)
    if not stage_output or not isinstance(stage_output.payload, dict):
        return {}
    return stage_output.payload


def download_fixture(request: HttpRequest, artifact: str) -> HttpResponse:
    document_model = _active_or_sample_document(request)
    document = _document_payload(document_model)
    outputs = _load_json_fixture("sample_outputs.json")
    _seed_stage_outputs_for_document(document_model, outputs)
    stage_outputs = _stage_outputs_for_document(document_model)
    stage_notes = list(
        ReviewNote.objects.filter(
            document=document_model,
            requested_action="chatbot_rerun_request",
        )
        .select_related("stage_output", "clause")
        .order_by("-created_at")
    )
    proposals = list(document_model.new_id_proposals.select_related("source_clause").all())
    proposal_rows = [_proposal_payload(proposal) for proposal in proposals]
    accepted_rows = [row for row in proposal_rows if row["status"] != NewIdProposal.Status.REJECTED]

    payloads: dict[str, Any] = {
        "summary": {
            "document": document["metadata"],
            "summary": _stage_payload_or_empty(stage_outputs, "summary_keywords").get("summary", {}),
            "keywords": _stage_payload_or_empty(stage_outputs, "summary_keywords").get("keywords", {}),
            "provenance": stage_outputs.get("summary_keywords").provenance if stage_outputs.get("summary_keywords") else {},
        },
        "entity-registry": {
            "document": document["metadata"],
            "entity_registry": _stage_payload_or_empty(stage_outputs, "entity_registry").get("entity_registry", []),
            "raw_output": stage_outputs.get("entity_registry").raw_output if stage_outputs.get("entity_registry") else "",
            "entity_review": _stage_payload_or_empty(stage_outputs, "entity_registry").get("entity_review", {}),
            "provenance": stage_outputs.get("entity_registry").provenance if stage_outputs.get("entity_registry") else {},
        },
        "key-events": {
            "document": document["metadata"],
            "occurrences_registry": _stage_payload_or_empty(stage_outputs, "occurrences_registry").get("occurrences_registry", []),
            "event_headword_review": _stage_payload_or_empty(stage_outputs, "occurrences_registry").get("event_headword_review", {}),
            "provenance": stage_outputs.get("occurrences_registry").provenance if stage_outputs.get("occurrences_registry") else {},
        },
    }
    payloads["final-review"] = build_workflow_export(document_model)

    if artifact == "new-ids":
        stream = io.StringIO()
        writer = csv.DictWriter(
            stream,
            fieldnames=[
                "proposed_id",
                "type",
                "headword",
                "form_evidence",
                "source_clause",
                "status",
                "reviewer_note",
            ],
        )
        writer.writeheader()
        for row in proposal_rows:
            writer.writerow(
                {
                    "proposed_id": row.get("proposed_id", ""),
                    "type": row.get("type", ""),
                    "headword": row.get("headword", ""),
                    "form_evidence": row.get("evidence_form", ""),
                    "source_clause": row.get("source_clause", ""),
                    "status": row.get("status", ""),
                    "reviewer_note": row.get("reviewer_note", ""),
                }
            )
        response = HttpResponse(stream.getvalue(), content_type="text/csv")
        response["Content-Disposition"] = 'attachment; filename="emtl_new_id_review_local.csv"'
        return response

    if artifact not in payloads:
        raise Http404("Unknown export")

    response = JsonResponse(payloads[artifact], json_dumps_params={"indent": 2})
    filename = (
        f"emtl_{payloads[artifact]['export_id']}.json"
        if artifact == "final-review"
        else f"emtl_{artifact}_local.json"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def data_page(request: HttpRequest) -> HttpResponse:
    return render(
        request,
        "tagger/info_page.html",
        {
            "nav_active": "data",
            "page_title": "Data",
            "eyebrow": "Reference layer",
            "intro": "Future home for the controlled vocabulary and database-backed review material that support EMTL tagging.",
            "sections": [
                ("Controlled vocabulary", "Headword-ID lists, type families, keywords, and authority records."),
                ("NEW-ID approvals", "Historian review queues for provisional IDs before they enter shared lookup tables."),
                ("Database support", "PostgreSQL lookup tables and checker/extractor support will plug in here later."),
            ],
        },
    )


def corpus_page(request: HttpRequest) -> HttpResponse:
    return render(
        request,
        "tagger/info_page.html",
        {
            "nav_active": "corpus",
            "page_title": "Corpus",
            "eyebrow": "Document layer",
            "intro": "Future home for processed documents, metadata, annotation layers, and corpus-scale review workflows.",
            "sections": [
                ("Processed documents", "A browsable list of uploaded, parsed, tagged, and reviewed source texts."),
                ("Annotation layers", "Clause-linked and hover-linked reading views for entities, occurrences, and full tags."),
                ("Search and correction", "Future corpus search, RAG-facing retrieval, and retroactive correction tools."),
            ],
        },
    )
